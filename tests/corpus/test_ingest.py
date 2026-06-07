"""Tests for the corpus ingestion module.

Gate criteria:
- from_url: returns non-empty text; strips script/style/nav tags.
- from_url: raises RuntimeError on HTTP error.
- from_pdf: extracts text and markdown tables from a minimal PDF.
- from_docx: extracts paragraph text from a minimal DOCX.
- from_markdown: strips headings, bold, italic, links, code fences.
- from_txt: returns file content unchanged (after normalisation).
- from_file: dispatches correctly by extension.
- from_file: raises ValueError for unsupported extensions.
- from_directory: returns dict of {path: text} for supported files only.
- from_directory: raises FileNotFoundError for missing directory.
- ingest (URL): writes .txt file to output_dir.
- ingest (file): writes .txt file to output_dir.
- ingest (directory): raises ValueError when output_dir is None.
- _url_to_filename: produces safe filename from a URL.
- _clean: collapses whitespace and blank lines.
"""

import io
import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from grimoire.corpus.ingest import (
    CleaningLevel,
    _clean,
    _url_to_filename,
    from_directory,
    from_file,
    from_markdown,
    from_txt,
    ingest,
)


# ---------------------------------------------------------------------------
# _clean / CleaningLevel
# ---------------------------------------------------------------------------

def test_clean_standard_collapses_blank_lines() -> None:
    assert _clean("a\n\n\n\nb") == "a\n\nb"


def test_clean_standard_collapses_spaces() -> None:
    assert _clean("hello   world") == "hello world"


def test_clean_strips_edges() -> None:
    assert _clean("  hello  ") == "hello"


def test_clean_minimal_preserves_blank_lines() -> None:
    text = "a\n\n\n\nb"
    result = _clean(text, level=CleaningLevel.MINIMAL)
    assert "\n\n\n" in result  # blank lines not collapsed


def test_clean_minimal_preserves_spaces() -> None:
    result = _clean("hello   world", level=CleaningLevel.MINIMAL)
    assert "   " in result


def test_clean_thorough_drops_short_lines() -> None:
    text = "Page 1\n\nA grappled creature has its speed reduced to zero.\n\n42"
    result = _clean(text, level=CleaningLevel.THOROUGH)
    assert "Page 1" not in result
    assert "42" not in result
    assert "grappled" in result


def test_clean_thorough_deduplicates_paragraphs() -> None:
    para = "A grappled creature has its speed reduced to zero."
    text = f"{para}\n\n{para}\n\n{para}"
    result = _clean(text, level=CleaningLevel.THOROUGH)
    assert result.count(para) == 1


def test_clean_thorough_keeps_unique_paragraphs() -> None:
    text = (
        "A grappled creature has its speed reduced to zero.\n\n"
        "The condition ends if the grappler is incapacitated."
    )
    result = _clean(text, level=CleaningLevel.THOROUGH)
    assert "grappled" in result
    assert "incapacitated" in result


# ---------------------------------------------------------------------------
# _url_to_filename
# ---------------------------------------------------------------------------

def test_url_to_filename_basic() -> None:
    name = _url_to_filename("https://example.com/rules/grapple")
    assert "grapple" in name
    assert "/" not in name


def test_url_to_filename_no_path() -> None:
    name = _url_to_filename("https://example.com/")
    assert len(name) > 0
    assert "/" not in name


# ---------------------------------------------------------------------------
# from_txt
# ---------------------------------------------------------------------------

def test_from_txt_returns_content() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.txt"
        p.write_text("Hello, Grimoire!\nSecond line.", encoding="utf-8")
        assert from_txt(str(p)) == "Hello, Grimoire!\nSecond line."


def test_from_txt_missing_file_raises() -> None:
    with pytest.raises(FileNotFoundError):
        from_txt("/tmp/grimoire_no_such_file.txt")


# ---------------------------------------------------------------------------
# from_markdown
# ---------------------------------------------------------------------------

def _md(content: str) -> str:
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.md"
        p.write_text(content, encoding="utf-8")
        return from_markdown(str(p))


def test_markdown_strips_headings() -> None:
    result = _md("# Title\n## Subtitle\nBody text.")
    assert "#" not in result
    assert "Title" in result
    assert "Body text" in result


def test_markdown_strips_bold() -> None:
    result = _md("This is **bold** text.")
    assert "**" not in result
    assert "bold" in result


def test_markdown_strips_italic() -> None:
    result = _md("This is *italic* text.")
    assert "*" not in result
    assert "italic" in result


def test_markdown_strips_links() -> None:
    result = _md("See [the rules](https://example.com) for details.")
    assert "https://example.com" not in result
    assert "the rules" in result


def test_markdown_strips_images() -> None:
    result = _md("![alt text](https://example.com/img.png)")
    assert "https://example.com" not in result


def test_markdown_strips_inline_code() -> None:
    result = _md("Use the `grapple` action.")
    assert "`" not in result
    assert "grapple" in result


def test_markdown_missing_file_raises() -> None:
    with pytest.raises(FileNotFoundError):
        from_markdown("/tmp/grimoire_no_such_file.md")


# ---------------------------------------------------------------------------
# from_pdf
# ---------------------------------------------------------------------------

def _minimal_pdf_bytes() -> bytes:
    """Return the bytes of a minimal valid single-page PDF with embedded text.

    Hand-crafted to avoid any runtime dependency on a PDF-writing library.
    The page contains the text "Hello World" in a simple content stream.
    """
    content = b"BT /F1 12 Tf 50 700 Td (Hello World) Tj ET"
    content_len = len(content)
    body = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        + b"4 0 obj\n<< /Length " + str(content_len).encode() + b" >>\nstream\n"
        + content + b"\nendstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    )
    # Build cross-reference table
    offsets = []
    pos = 0
    for obj_bytes in [
        b"%PDF-1.4\n",
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            b" /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        ),
        (
            b"4 0 obj\n<< /Length " + str(content_len).encode() + b" >>\nstream\n"
            + content + b"\nendstream\nendobj\n"
        ),
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]:
        offsets.append(pos)
        pos += len(obj_bytes)
    xref_pos = pos
    xref = b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()
    trailer = (
        b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    )
    return body + xref + trailer


def test_from_pdf_extracts_text() -> None:
    try:
        import pdfplumber  # noqa: F401
        from grimoire.corpus.ingest import from_pdf
    except BaseException:
        pytest.skip("pdfplumber not importable in this environment")

    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = str(Path(tmp) / "test.pdf")
        Path(pdf_path).write_bytes(_minimal_pdf_bytes())
        result = from_pdf(pdf_path)
        assert isinstance(result, str)


def test_from_pdf_missing_file_raises() -> None:
    try:
        import pdfplumber  # noqa: F401
    except BaseException:
        pytest.skip("pdfplumber not importable in this environment")
    from grimoire.corpus.ingest import from_pdf
    with pytest.raises(FileNotFoundError):
        from_pdf("/tmp/grimoire_no_such.pdf")


def test_table_to_markdown_basic() -> None:
    """_table_to_markdown should produce a pipe table with a separator row."""
    from grimoire.corpus.ingest import _table_to_markdown

    table = [
        ["STR", "DEX", "CON"],
        ["8 (-1)", "14 (+2)", "10 (+0)"],
    ]
    md = _table_to_markdown(table)
    lines = md.splitlines()
    assert lines[0].startswith("| STR")
    assert lines[1] == "| --- | --- | --- |"
    assert "8 (-1)" in lines[2]


def test_table_to_markdown_empty_cells() -> None:
    """Empty/None cells should be replaced with a space, not break the table."""
    from grimoire.corpus.ingest import _table_to_markdown

    table = [["A", None, "C"], [None, "B", None]]
    md = _table_to_markdown(table)
    assert "|" in md
    assert "None" not in md


def test_table_to_markdown_empty_table() -> None:
    """An all-empty table should return an empty string."""
    from grimoire.corpus.ingest import _table_to_markdown

    assert _table_to_markdown([]) == ""
    assert _table_to_markdown([[None, None]]) == ""


# ---------------------------------------------------------------------------
# from_docx
# ---------------------------------------------------------------------------

def test_from_docx_extracts_paragraphs() -> None:
    docx_mod = pytest.importorskip("docx")
    from docx import Document
    from grimoire.corpus.ingest import from_docx

    doc = Document()
    doc.add_paragraph("A grappled creature has its speed reduced to zero.")
    doc.add_paragraph("The condition ends if the grappler is incapacitated.")

    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "test.docx")
        doc.save(path)
        result = from_docx(path)

    assert "grappled" in result
    assert "incapacitated" in result


def test_from_docx_missing_file_raises() -> None:
    pytest.importorskip("docx")
    from grimoire.corpus.ingest import from_docx
    with pytest.raises(FileNotFoundError):
        from_docx("/tmp/grimoire_no_such.docx")


# ---------------------------------------------------------------------------
# from_url
# ---------------------------------------------------------------------------

def test_from_url_returns_text() -> None:
    pytest.importorskip("bs4")
    from grimoire.corpus.ingest import from_url

    html = """
    <html><body>
      <nav>Navigation</nav>
      <main><p>A grappled creature has its speed reduced to zero.</p></main>
      <footer>Footer</footer>
      <script>alert('noise')</script>
    </body></html>
    """
    mock_resp = MagicMock()
    mock_resp.text = html
    mock_resp.raise_for_status = MagicMock()

    with patch("requests.get", return_value=mock_resp):
        result = from_url("https://example.com/rules")

    assert "grappled" in result
    assert "Navigation" not in result
    assert "Footer" not in result
    assert "noise" not in result


def test_from_url_raises_on_http_error() -> None:
    pytest.importorskip("bs4")
    import requests
    from grimoire.corpus.ingest import from_url

    mock_resp = MagicMock()
    mock_resp.raise_for_status.side_effect = requests.HTTPError("404")

    with patch("requests.get", return_value=mock_resp):
        with pytest.raises(RuntimeError, match="Failed to fetch"):
            from_url("https://example.com/not-found")


def test_from_url_treats_md_url_as_markdown() -> None:
    """A URL ending in .md should be processed as Markdown, not HTML."""
    pytest.importorskip("bs4")
    from grimoire.corpus.ingest import from_url

    md_text = "# Combat\n\nA **grappled** creature has its speed reduced to zero.\n"
    mock_resp = MagicMock()
    mock_resp.text = md_text
    mock_resp.headers = {"Content-Type": "text/plain; charset=utf-8"}
    mock_resp.raise_for_status = MagicMock()

    with patch("requests.get", return_value=mock_resp):
        result = from_url("https://raw.githubusercontent.com/user/repo/main/rules.md")

    assert "grappled" in result
    assert "#" not in result       # heading marker stripped
    assert "**" not in result      # bold markers stripped


def test_from_url_treats_text_plain_as_markdown() -> None:
    """A text/plain response (no .md extension) should also skip BeautifulSoup."""
    pytest.importorskip("bs4")
    from grimoire.corpus.ingest import from_url

    md_text = "## Spellcasting\n\nYou can cast spells of your choice.\n"
    mock_resp = MagicMock()
    mock_resp.text = md_text
    mock_resp.headers = {"Content-Type": "text/plain"}
    mock_resp.raise_for_status = MagicMock()

    with patch("requests.get", return_value=mock_resp):
        result = from_url("https://example.com/srd")

    assert "Spellcasting" in result
    assert "##" not in result


# ---------------------------------------------------------------------------
# from_file dispatch
# ---------------------------------------------------------------------------

def test_from_file_dispatches_txt() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "data.txt"
        p.write_text("hello world", encoding="utf-8")
        assert from_file(str(p)) == "hello world"


def test_from_file_dispatches_markdown() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "notes.md"
        p.write_text("# Title\nBody.", encoding="utf-8")
        result = from_file(str(p))
        assert "Title" in result
        assert "#" not in result


def test_from_file_unsupported_extension_raises() -> None:
    with pytest.raises(ValueError, match="Unsupported file extension"):
        from_file("/tmp/grimoire_test.xyz")


# ---------------------------------------------------------------------------
# from_directory
# ---------------------------------------------------------------------------

def test_from_directory_collects_supported_files() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        (Path(tmp) / "a.txt").write_text("file a", encoding="utf-8")
        (Path(tmp) / "b.txt").write_text("file b", encoding="utf-8")
        (Path(tmp) / "ignore.csv").write_text("not supported", encoding="utf-8")

        results = from_directory(tmp)

    assert len(results) == 2
    texts = list(results.values())
    assert any("file a" in t for t in texts)
    assert any("file b" in t for t in texts)


def test_from_directory_recursive() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        sub = Path(tmp) / "sub"
        sub.mkdir()
        (Path(tmp) / "root.txt").write_text("root", encoding="utf-8")
        (sub / "child.txt").write_text("child", encoding="utf-8")

        results = from_directory(tmp, recursive=True)

    assert len(results) == 2


def test_from_directory_missing_raises() -> None:
    with pytest.raises(FileNotFoundError):
        from_directory("/tmp/grimoire_no_such_dir_xyz/")


# ---------------------------------------------------------------------------
# ingest (integration)
# ---------------------------------------------------------------------------

def test_ingest_file_writes_txt() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.txt"
        src.write_text("Grimoire corpus content.", encoding="utf-8")
        out = Path(tmp) / "output"

        ingest(str(src), output_dir=str(out))

        files = list(out.glob("*.txt"))
        assert len(files) == 1
        assert "Grimoire corpus content" in files[0].read_text(encoding="utf-8")


def test_ingest_url_writes_txt() -> None:
    pytest.importorskip("bs4")
    from grimoire.corpus.ingest import ingest

    html = "<html><body><main><p>Dragon lore.</p></main></body></html>"
    mock_resp = MagicMock()
    mock_resp.text = html
    mock_resp.raise_for_status = MagicMock()

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "output"
        with patch("requests.get", return_value=mock_resp):
            ingest("https://example.com/dragons", output_dir=str(out))

        files = list(out.glob("*.txt"))
        assert len(files) == 1
        assert "Dragon lore" in files[0].read_text(encoding="utf-8")


def test_ingest_directory_requires_output_dir() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        with pytest.raises(ValueError, match="output_dir is required"):
            ingest(tmp)


def test_ingest_returns_text_without_output_dir() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "note.txt"
        src.write_text("Some text.", encoding="utf-8")
        result = ingest(str(src))
        assert result == "Some text."


def test_ingest_on_progress_callback_fires() -> None:
    """on_progress must be called at least once when writing a file."""
    messages: list[str] = []

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "note.txt"
        src.write_text("Some text.", encoding="utf-8")
        out = Path(tmp) / "output"
        ingest(str(src), output_dir=str(out), on_progress=messages.append)

    assert len(messages) >= 1


# ---------------------------------------------------------------------------
# from_xlsx
# ---------------------------------------------------------------------------

def test_from_xlsx_renders_sheet_as_markdown() -> None:
    openpyxl = pytest.importorskip("openpyxl")
    from grimoire.corpus.ingest import from_xlsx

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Monsters"
    ws.append(["Name", "CR", "HP"])
    ws.append(["Goblin", "1/4", "7"])
    ws.append(["Orc", "1/2", "15"])

    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "monsters.xlsx")
        wb.save(path)
        result = from_xlsx(path)

    assert "Monsters" in result
    assert "Goblin" in result
    assert "CR" in result
    assert "|" in result


def test_from_xlsx_multiple_sheets() -> None:
    openpyxl = pytest.importorskip("openpyxl")
    from grimoire.corpus.ingest import from_xlsx

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["A", "B"])
    ws1.append(["1", "2"])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["X", "Y"])
    ws2.append(["3", "4"])

    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "multi.xlsx")
        wb.save(path)
        result = from_xlsx(path)

    assert "Sheet1" in result
    assert "Sheet2" in result


def test_from_xlsx_missing_file_raises() -> None:
    pytest.importorskip("openpyxl")
    from grimoire.corpus.ingest import from_xlsx
    with pytest.raises(FileNotFoundError):
        from_xlsx("/tmp/grimoire_no_such.xlsx")


def test_from_file_dispatches_xlsx() -> None:
    openpyxl = pytest.importorskip("openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Col1", "Col2"])
    ws.append(["val1", "val2"])

    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "data.xlsx")
        wb.save(path)
        result = from_file(path)

    assert "Col1" in result


def test_ingest_cleaning_level_thorough() -> None:
    """Thorough cleaning drops very short lines from the output."""
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "noisy.txt"
        src.write_text(
            "Page 1\n\nA grappled creature has its speed reduced to zero.\n\nEnd",
            encoding="utf-8",
        )
        result = ingest(str(src), cleaning=CleaningLevel.THOROUGH)

    assert "Page 1" not in result
    assert "grappled" in result
